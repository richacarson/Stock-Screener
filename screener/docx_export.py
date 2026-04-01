"""Generate Word (.docx) reports matching the IOWN template format."""

import io
from pathlib import Path

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Cache logo bytes
_logo_bytes: bytes | None = None
LOGO_URL = "https://richacarson.github.io/Dashboard/iown-logo.png"

# Colors
C_DARK = RGBColor(0x1A, 0x20, 0x10)
C_T2 = RGBColor(0x3A, 0x4A, 0x28)
C_T3 = RGBColor(0x6E, 0x84, 0x50)
C_MUTED = RGBColor(0x9D, 0xAF, 0x88)
C_ACCENT = RGBColor(0x4A, 0x6B, 0x25)
C_GREEN = RGBColor(0x16, 0xA3, 0x4A)
C_GOLD = RGBColor(0xD9, 0x77, 0x06)
C_DARKGOLD = RGBColor(0xB8, 0x86, 0x0B)
C_RED = RGBColor(0xDC, 0x26, 0x26)
C_BLUE = RGBColor(0x25, 0x63, 0xEB)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

FOOTER_TEXT = (
    'Paradiem, LLC dba Intentional Ownership (\u201CIOWN\u201D) is a registered '
    'investment advisor. The information provided is for educational and '
    'informational purposes only and does not constitute investment advice.'
)


def _get_logo() -> bytes | None:
    global _logo_bytes
    if _logo_bytes is not None:
        return _logo_bytes
    logo_path = Path("templates/iown-logo.png")
    if logo_path.exists():
        _logo_bytes = logo_path.read_bytes()
        return _logo_bytes
    try:
        resp = requests.get(LOGO_URL, timeout=10)
        resp.raise_for_status()
        _logo_bytes = resp.content
        logo_path.write_bytes(_logo_bytes)
        return _logo_bytes
    except Exception:
        return None


def _run(paragraph, text: str, size: Pt, bold: bool = False, color=C_DARK):
    run = paragraph.add_run(text)
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return run


def _rec_color(rec: str) -> RGBColor:
    return {"BUY": C_GREEN, "HOLD": C_GOLD, "SELL": C_RED}.get(
        rec.upper(), C_BLUE
    )


def _rec_bg(rec: str) -> str:
    """Return hex background color for recommendation badge."""
    return {
        "BUY": "E8F5E9", "HOLD": "FFF8E1", "SELL": "FFEBEE", "WATCH": "E3F2FD"
    }.get(rec.upper(), "F5F5F5")


def _score_color(score: int) -> RGBColor:
    if score >= 7:
        return C_GREEN
    if score >= 4:
        return C_DARKGOLD
    return C_RED


def _score_bar_hex(score: int) -> str:
    if score >= 7:
        return "16A34A"
    if score >= 4:
        return "B8860B"
    return "DC2626"


def _label_color(label: str) -> RGBColor:
    label = label.upper()
    if label in ("STRONG", "SAFE", "LOW RISK", "INFINITE"):
        return C_GREEN
    if label in ("WEAK", "AT RISK", "HIGH RISK", "FINITE"):
        return C_RED
    return C_DARKGOLD


def _set_cell_shading(cell, hex_color: str):
    """Set background color on a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border(cell, side: str, color: str, size: str = "6"):
    """Set a single border on a cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tc_pr.append(borders)
    border = parse_xml(
        f'<w:{side} {nsdecls("w")} w:val="single" w:sz="{size}" '
        f'w:space="0" w:color="{color}"/>'
    )
    existing = borders.find(qn(f"w:{side}"))
    if existing is not None:
        borders.remove(existing)
    borders.append(border)


def _remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}/>'
    )
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.append(borders)


def _add_score_bar(doc: Document, score: int, max_score: int = 10):
    """Add a visual score bar using a 2-cell table (filled | empty)."""
    pct = min(max(score / max_score, 0), 1)
    fill_w = max(int(pct * 100), 1)
    empty_w = 100 - fill_w

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_table_borders(table)

    # Set overall table width
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = parse_xml(f'<w:tblW {nsdecls("w")} w:w="3500" w:type="dxa"/>')
    existing_w = tbl_pr.find(qn("w:tblW"))
    if existing_w is not None:
        tbl_pr.remove(existing_w)
    tbl_pr.append(tbl_w)

    row = table.rows[0]
    row.height = Cm(0.3)

    # Filled cell
    fill_cell = row.cells[0]
    fill_cell.width = Emu(int(fill_w * 35000))
    _set_cell_shading(fill_cell, _score_bar_hex(score))
    fill_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    fill_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    _run(fill_cell.paragraphs[0], "", Pt(2))

    # Empty cell (track)
    empty_cell = row.cells[1]
    empty_cell.width = Emu(int(empty_w * 35000))
    _set_cell_shading(empty_cell, "EDEFE8")
    empty_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    empty_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    _run(empty_cell.paragraphs[0], "", Pt(2))


def _section_header(doc: Document, text: str, border_color: str = "B8860B"):
    """Add a section header with colored bottom border."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_table_borders(table)

    # Full width
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    existing_w = tbl_pr.find(qn("w:tblW"))
    if existing_w is not None:
        tbl_pr.remove(existing_w)
    tbl_pr.append(tbl_w)

    cell = table.rows[0].cells[0]
    _set_cell_border(cell, "bottom", border_color, "8")

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = _run(p, text.upper(), Pt(8.5), bold=True, color=C_T3)
    r.font.letter_spacing = Pt(1)


def _score_block(doc: Document, label: str, tag: str, score, analysis: str):
    """Add a dimension score with label, colored tag, score bar, and analysis."""
    score_val = score if isinstance(score, int) else 0

    # Label + tag line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    _run(p, label, Pt(11), bold=True, color=C_DARK)
    if tag:
        _run(p, "  ", Pt(11))
        r = _run(p, f" {tag} ", Pt(8), bold=True, color=_label_color(tag))

    # Score number + /10
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    _run(p2, str(score_val), Pt(14), bold=True, color=_score_color(score_val))
    _run(p2, " /10", Pt(9), color=C_MUTED)

    # Score bar
    _add_score_bar(doc, score_val)

    # Analysis text
    if analysis:
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_before = Pt(3)
        p3.paragraph_format.space_after = Pt(6)
        _run(p3, analysis, Pt(9), color=C_T2)


def _add_rec_badge(doc: Document, rec: str):
    """Add a colored recommendation badge."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_table_borders(table)

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = parse_xml(f'<w:tblW {nsdecls("w")} w:w="1200" w:type="dxa"/>')
    existing_w = tbl_pr.find(qn("w:tblW"))
    if existing_w is not None:
        tbl_pr.remove(existing_w)
    tbl_pr.append(tbl_w)

    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, _rec_bg(rec))
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    _run(p, rec, Pt(10), bold=True, color=_rec_color(rec))


def _add_attr_tags(doc: Document, label_text: str, attrs: list, is_negative: bool):
    """Add colored attribution tags inline."""
    if not attrs:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    color = C_RED if is_negative else C_GREEN
    _run(p, f"{label_text}: ", Pt(9), bold=True, color=C_T3)
    for i, attr in enumerate(attrs):
        _run(p, f" {attr} ", Pt(8.5), bold=True, color=color)
        if i < len(attrs) - 1:
            _run(p, "  \u00b7  ", Pt(8.5), color=C_MUTED)


def _bullet(doc: Document, text: str, number: int | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    prefix = f"{number}. " if number else "\u2022  "
    _run(p, prefix, Pt(10), bold=True, color=C_ACCENT)
    _run(p, text, Pt(10), color=C_T2)


def generate_docx(report: dict) -> bytes:
    """Generate a Word doc from a report dict, return bytes."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Logo
    logo = _get_logo()
    if logo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(io.BytesIO(logo), width=Inches(1.5))

    ticker = report.get("ticker", "")
    name = report.get("name", "")
    sleeve = report.get("sleeve", "")
    rec = report.get("recommendation", "")
    score = report.get("overall_score", 0)
    screen_date = report.get("screen_date", "")
    profile = report.get("profile", {})
    excellence = report.get("excellence_evaluation", {})
    ai_res = report.get("ai_resilience", {})
    ig = report.get("infinite_game", {})
    faith = report.get("faith_alignment", {})

    inspire_score = faith.get("inspire_impact_score", "N/A")
    mindset = ig.get("mindset", "")

    # ── Title ──
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    _run(p, ticker, Pt(28), bold=True, color=C_ACCENT)
    _run(p, f"  {name}", Pt(20), bold=True, color=C_DARK)

    # Subtitle metadata
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    parts = [sleeve.upper() + " SLEEVE", screen_date]
    if inspire_score is not None and inspire_score != "N/A":
        parts.append(f"Inspire: {inspire_score}")
    if mindset:
        parts.append(mindset)
    _run(p, " \u00b7 ".join(parts), Pt(9), bold=True, color=C_T3)

    # Recommendation badge
    _add_rec_badge(doc, rec)

    # Overall score
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    _run(p, str(score), Pt(40), bold=True, color=C_DARK)
    _run(p, " / ", Pt(14), color=C_MUTED)
    _run(p, "100", Pt(12), color=C_MUTED)

    # ── Company Profile ──
    _section_header(doc, "Company Profile", "B8860B")
    sector = profile.get("sector", "")
    industry = profile.get("industry", "")
    exchange = profile.get("exchange", "")
    country = profile.get("country", "")
    meta_parts = [x for x in [sector, industry, exchange, country] if x]
    if meta_parts:
        p = doc.add_paragraph()
        _run(p, " \u00b7 ".join(meta_parts), Pt(9), bold=True, color=C_T3)

    desc = profile.get("description", "")
    if desc:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, desc, Pt(10), color=C_T2)

    employees = profile.get("employees")
    website = profile.get("website", "")
    info_parts = []
    if employees:
        info_parts.append(f"{employees:,} Employees")
    if website:
        info_parts.append(website)
    if info_parts:
        p = doc.add_paragraph()
        _run(p, " \u00b7 ".join(info_parts), Pt(9), bold=True, color=C_T3)

    # ── Excellence Evaluation ──
    _section_header(doc, "Excellence Evaluation \u2014 Think Like an Owner (50%)", "16A34A")

    for dim_key, dim_label in [("innovation", "Innovation"),
                                ("inspiration", "Inspiration"),
                                ("infrastructure", "Infrastructure")]:
        d = excellence.get(dim_key, {})
        _score_block(doc, dim_label, d.get("label", ""),
                     d.get("score", 0), d.get("analysis", ""))

    # ── Infinite Game ──
    _section_header(doc, "Finite vs Infinite Game \u2014 Sinek (25%)", "16A34A")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _run(p, "Mindset: ", Pt(11), color=C_T3)
    _run(p, mindset, Pt(11), bold=True, color=_label_color(mindset))
    _run(p, "      Overall: ", Pt(11), color=C_T3)
    _run(p, str(ig.get("overall", 0)), Pt(14), bold=True,
         color=_score_color(ig.get("overall", 0)))
    _run(p, " /10", Pt(9), color=C_MUTED)

    summary = ig.get("summary", "")
    if summary:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _run(p, summary, Pt(9), color=C_T2)

    for dim in ["just_cause", "trusting_teams", "worthy_rivals",
                "existential_flexibility", "courage_to_lead"]:
        d = ig.get(dim, {})
        label = dim.replace("_", " ").title()
        _score_block(doc, label, "", d.get("score", 0), d.get("analysis", ""))

    # ── Investment Thesis ──
    _section_header(doc, "Investment Thesis", "B8860B")
    thesis = report.get("investment_thesis", "")
    if thesis:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _run(p, thesis, Pt(10), color=C_T2)
    thesis2 = report.get("thesis_continued", "")
    if thesis2:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _run(p, thesis2, Pt(10), color=C_T2)

    # ── Key Catalysts ──
    catalysts = report.get("key_catalysts", [])
    if catalysts:
        _section_header(doc, "Key Catalysts", "16A34A")
        for i, c in enumerate(catalysts, 1):
            _bullet(doc, c, i)

    # ── Key Risks ──
    risks = report.get("key_risks", [])
    if risks:
        _section_header(doc, "Key Risks", "DC2626")
        for i, r in enumerate(risks, 1):
            _bullet(doc, r, i)

    # ── AI Resilience ──
    _section_header(doc, "AI Resilience (25%)", "B8860B")
    _score_block(doc, "AI Resilience", ai_res.get("label", ""),
                 ai_res.get("score", 0), ai_res.get("analysis", ""))

    # ── Faith Alignment ──
    _section_header(doc, "Faith Alignment \u2014 Inspire Insight", "B8860B")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _run(p, "Inspire Impact Score:  ", Pt(11), color=C_T3)
    isc = inspire_score if inspire_score is not None else "N/A"
    isc_color = C_RED if isinstance(inspire_score, (int, float)) and inspire_score < 0 else C_GREEN
    _run(p, str(isc), Pt(14), bold=True, color=isc_color)
    fl = faith.get("label", "")
    if fl:
        _run(p, f"    {fl}", Pt(10), bold=True, color=_label_color(fl))

    neg = faith.get("negative_attributions", [])
    _add_attr_tags(doc, "Negative", neg, True)

    pos = faith.get("positive_attributions", [])
    _add_attr_tags(doc, "Positive", pos, False)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    _run(p, "Source: Inspire Insight", Pt(7.5), color=C_MUTED)

    # ── Resources ──
    sources = report.get("sources", [])
    if sources:
        _section_header(doc, "Resources", "B8860B")
        for i, s in enumerate(sources, 1):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(0.3)
            _run(p, f"{i}. ", Pt(7.5), bold=True, color=C_ACCENT)
            _run(p, s, Pt(8.5), color=C_T2)

    # ── Footer ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    _run(p, FOOTER_TEXT, Pt(8), color=C_MUTED)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
